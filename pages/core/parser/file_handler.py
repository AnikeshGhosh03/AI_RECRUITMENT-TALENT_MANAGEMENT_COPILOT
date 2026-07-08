import os
from pathlib import Path
from typing import BinaryIO, Optional

import docx
import fitz


def extract_text_from_path(file_path: str | os.PathLike) -> str:
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        return extract_text_from_pdf(path)
    if suffix == ".docx":
        return extract_text_from_docx(path)
    if suffix == ".txt":
        return path.read_text(encoding="utf-8", errors="ignore")
    raise ValueError(f"Unsupported file type: {suffix}")


def extract_text_from_pdf(file_path: str | os.PathLike) -> str:
    """Extract text and hyperlinks from PDF."""
    document = fitz.open(file_path)
    text_chunks = []
    
    for page in document:
        # Get regular text
        text = page.get_text()
        if text and text.strip():
            text_chunks.append(text.strip())
        
        # Extract hyperlinks
        links = page.get_links()
        for link in links:
            if "uri" in link and link["uri"]:
                # Append the URL to the text so it can be extracted
                text_chunks.append(f"\n{link['uri']}")
    
    document.close()
    return "\n".join(text_chunks)


def extract_text_from_docx(file_path: str | os.PathLike) -> str:
    """Extract text and hyperlinks from DOCX."""
    document = docx.Document(file_path)
    text_chunks = []
    
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            text_chunks.append(paragraph.text)
        
        # Extract hyperlinks from runs
        for run in paragraph.runs:
            if run.element.rPr is not None:
                # Check for hyperlink style
                hyperlinks = run.element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hyperlink', 
                                                 run.element.nsmap if hasattr(run.element, 'nsmap') else None)
                for hyperlink in hyperlinks:
                    rel_id = hyperlink.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                    if rel_id:
                        try:
                            url = document.part.rels[rel_id].target_ref
                            if url:
                                text_chunks.append(f"\n{url}")
                        except (KeyError, AttributeError):
                            pass
    
    # Also check hyperlinks at paragraph level
    for rel in document.part.rels.values():
        if "hyperlink" in rel.reltype:
            text_chunks.append(f"\n{rel.target_ref}")
    
    return "\n".join(text_chunks)


def extract_text_from_uploaded_file(uploaded_file: BinaryIO, filename: Optional[str] = None) -> str:
    suffix = Path(filename or "").suffix.lower()
    if suffix == ".pdf":
        return extract_text_from_bytes(uploaded_file.read(), suffix=".pdf")
    if suffix == ".docx":
        return extract_text_from_bytes(uploaded_file.read(), suffix=".docx")
    if suffix == ".txt":
        return uploaded_file.read().decode("utf-8", errors="ignore")
    raise ValueError(f"Unsupported file type: {suffix}")


def extract_text_from_bytes(file_bytes: bytes, suffix: str) -> str:
    """Extract text and hyperlinks from bytes."""
    if suffix == ".pdf":
        document = fitz.open(stream=file_bytes, filetype="pdf")
        text_chunks = []
        
        for page in document:
            # Get regular text
            text = page.get_text()
            if text and text.strip():
                text_chunks.append(text.strip())
            
            # Extract hyperlinks
            links = page.get_links()
            for link in links:
                if "uri" in link and link["uri"]:
                    text_chunks.append(f"\n{link['uri']}")
        
        document.close()
        return "\n".join(text_chunks)
    
    if suffix == ".docx":
        import io

        document = docx.Document(io.BytesIO(file_bytes))
        text_chunks = []
        
        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                text_chunks.append(paragraph.text)
            
            # Extract hyperlinks from runs
            for run in paragraph.runs:
                if run.element.rPr is not None:
                    hyperlinks = run.element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hyperlink',
                                                     run.element.nsmap if hasattr(run.element, 'nsmap') else None)
                    for hyperlink in hyperlinks:
                        rel_id = hyperlink.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                        if rel_id:
                            try:
                                url = document.part.rels[rel_id].target_ref
                                if url:
                                    text_chunks.append(f"\n{url}")
                            except (KeyError, AttributeError):
                                pass
        
        # Also check hyperlinks at paragraph level
        for rel in document.part.rels.values():
            if "hyperlink" in rel.reltype:
                text_chunks.append(f"\n{rel.target_ref}")
        
        return "\n".join(text_chunks)
    
    raise ValueError(f"Unsupported file type: {suffix}")
