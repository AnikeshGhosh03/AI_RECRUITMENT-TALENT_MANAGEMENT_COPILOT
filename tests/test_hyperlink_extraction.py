"""
Test hyperlink extraction from resume files.

This test verifies that LinkedIn, GitHub, and portfolio links are correctly
extracted from both PDF and DOCX files, even when they are embedded as hyperlinks
rather than visible text.
"""
import io
import os
from pathlib import Path

import docx
import fitz
import pytest

from pages.core.parser.field_extractor import (
    _extract_github,
    _extract_linkedin,
    _extract_portfolio,
    extract_candidate_profile,
)
from pages.core.parser.file_handler import (
    extract_text_from_bytes,
    extract_text_from_docx,
    extract_text_from_pdf,
)


class TestHyperlinkExtraction:
    """Test extraction of hyperlinked URLs from resumes."""

    def test_extract_linkedin_from_text(self):
        """Test LinkedIn URL extraction from various formats."""
        test_cases = [
            ("https://www.linkedin.com/in/johndoe", "https://www.linkedin.com/in/johndoe"),
            ("linkedin.com/in/johndoe", "https://www.linkedin.com/in/johndoe"),
            ("www.linkedin.com/in/johndoe", "https://www.linkedin.com/in/johndoe"),
            ("Contact: https://linkedin.com/in/jane-smith", "https://www.linkedin.com/in/jane-smith"),
        ]
        
        for text, expected in test_cases:
            result = _extract_linkedin(text)
            assert result == expected, f"Failed for: {text}"

    def test_extract_github_from_text(self):
        """Test GitHub URL extraction from various formats."""
        test_cases = [
            ("https://github.com/johndoe", "https://github.com/johndoe"),
            ("github.com/johndoe", "https://github.com/johndoe"),
            ("www.github.com/jane-smith", "https://www.github.com/jane-smith"),
            ("Check my work: https://github.com/developer123", "https://github.com/developer123"),
        ]
        
        for text, expected in test_cases:
            result = _extract_github(text)
            assert result == expected, f"Failed for: {text}"

    def test_extract_portfolio_from_text(self):
        """Test portfolio URL extraction while excluding social media."""
        test_cases = [
            ("https://johndoe.dev", "https://johndoe.dev"),
            ("www.myportfolio.io", "https://www.myportfolio.io"),
            ("mysite.com", "https://mysite.com"),
        ]
        
        for text, expected in test_cases:
            result = _extract_portfolio(text)
            assert result == expected, f"Failed for: {text}"
    
    def test_portfolio_excludes_social_media(self):
        """Test that portfolio extraction excludes social media sites."""
        social_urls = [
            "https://linkedin.com/in/test",
            "https://github.com/test",
            "https://facebook.com/test",
            "https://twitter.com/test",
        ]
        
        for url in social_urls:
            result = _extract_portfolio(url)
            assert result is None, f"Should not extract social media URL: {url}"

    def test_pdf_hyperlink_extraction(self, tmp_path):
        """Test that hyperlinks are extracted from PDF files."""
        # Create a test PDF with hyperlinks
        pdf_path = tmp_path / "test_resume.pdf"
        doc = fitz.open()
        page = doc.new_page()
        
        # Add visible text
        page.insert_text((50, 50), "John Doe")
        page.insert_text((50, 70), "Software Engineer")
        page.insert_text((50, 90), "Email: john@example.com")
        
        # Add hyperlinked text (text with embedded URL)
        page.insert_text((50, 110), "LinkedIn")
        rect = fitz.Rect(50, 100, 150, 120)
        page.insert_link({
            "kind": fitz.LINK_URI,
            "from": rect,
            "uri": "https://www.linkedin.com/in/johndoe"
        })
        
        page.insert_text((50, 130), "GitHub")
        rect2 = fitz.Rect(50, 120, 150, 140)
        page.insert_link({
            "kind": fitz.LINK_URI,
            "from": rect2,
            "uri": "https://github.com/johndoe"
        })
        
        doc.save(pdf_path)
        doc.close()
        
        # Extract text and verify hyperlinks are included
        extracted_text = extract_text_from_pdf(pdf_path)
        
        assert "john@example.com" in extracted_text
        assert "https://www.linkedin.com/in/johndoe" in extracted_text
        assert "https://github.com/johndoe" in extracted_text
        
        # Test full profile extraction
        profile = extract_candidate_profile(extracted_text, str(pdf_path))
        assert profile.full_name == "John Doe"
        assert profile.contact_info.email == "john@example.com"
        assert profile.contact_info.linkedin == "https://www.linkedin.com/in/johndoe"
        assert profile.contact_info.github == "https://github.com/johndoe"

    def test_docx_hyperlink_extraction(self, tmp_path):
        """Test that hyperlinks are extracted from DOCX files."""
        # Create a test DOCX with hyperlinks
        docx_path = tmp_path / "test_resume.docx"
        doc = docx.Document()
        
        # Add regular paragraphs
        doc.add_paragraph("Jane Smith")
        doc.add_paragraph("Data Scientist")
        doc.add_paragraph("Email: jane@example.com")
        
        # Add hyperlinked text
        p1 = doc.add_paragraph()
        p1.add_run("LinkedIn: ")
        add_hyperlink(p1, "https://www.linkedin.com/in/janesmith", "Profile")
        
        p2 = doc.add_paragraph()
        p2.add_run("GitHub: ")
        add_hyperlink(p2, "https://github.com/janesmith", "Code")
        
        p3 = doc.add_paragraph()
        p3.add_run("Portfolio: ")
        add_hyperlink(p3, "https://janesmith.dev", "Website")
        
        doc.save(docx_path)
        
        # Extract text and verify hyperlinks are included
        extracted_text = extract_text_from_docx(docx_path)
        
        assert "jane@example.com" in extracted_text
        assert "https://www.linkedin.com/in/janesmith" in extracted_text
        assert "https://github.com/janesmith" in extracted_text
        assert "https://janesmith.dev" in extracted_text
        
        # Test full profile extraction
        profile = extract_candidate_profile(extracted_text, str(docx_path))
        assert profile.full_name == "Jane Smith"
        assert profile.contact_info.email == "jane@example.com"
        assert profile.contact_info.linkedin == "https://www.linkedin.com/in/janesmith"
        assert profile.contact_info.github == "https://github.com/janesmith"
        assert profile.contact_info.portfolio == "https://janesmith.dev"

    def test_extract_from_bytes_pdf(self):
        """Test hyperlink extraction from PDF bytes (uploaded file scenario)."""
        # Create a PDF in memory
        doc = fitz.open()
        page = doc.new_page()
        
        page.insert_text((50, 50), "Test User")
        page.insert_text((50, 70), "test@example.com")
        
        # Add hyperlink
        rect = fitz.Rect(50, 80, 150, 100)
        page.insert_link({
            "kind": fitz.LINK_URI,
            "from": rect,
            "uri": "https://github.com/testuser"
        })
        
        # Save to bytes
        pdf_bytes = doc.tobytes()
        doc.close()
        
        # Extract and verify
        extracted_text = extract_text_from_bytes(pdf_bytes, ".pdf")
        assert "https://github.com/testuser" in extracted_text
        assert "test@example.com" in extracted_text

    def test_extract_from_bytes_docx(self):
        """Test hyperlink extraction from DOCX bytes (uploaded file scenario)."""
        doc = docx.Document()
        doc.add_paragraph("Test User")
        doc.add_paragraph("test@example.com")
        
        p = doc.add_paragraph()
        add_hyperlink(p, "https://github.com/testuser", "GitHub")
        
        # Save to bytes
        bio = io.BytesIO()
        doc.save(bio)
        docx_bytes = bio.getvalue()
        
        # Extract and verify
        extracted_text = extract_text_from_bytes(docx_bytes, ".docx")
        assert "https://github.com/testuser" in extracted_text
        assert "test@example.com" in extracted_text


def add_hyperlink(paragraph, url, text):
    """
    Helper function to add a hyperlink to a paragraph in python-docx.
    
    Note: python-docx doesn't have built-in hyperlink support, so we use XML manipulation.
    """
    # Get the paragraph's part
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    
    # Create the hyperlink element
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id)
    
    # Create a new run for the hyperlink text
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')
    
    # Add hyperlink styling (blue and underlined)
    color = docx.oxml.shared.OxmlElement('w:color')
    color.set(docx.oxml.shared.qn('w:val'), '0000FF')
    rPr.append(color)
    
    u = docx.oxml.shared.OxmlElement('w:u')
    u.set(docx.oxml.shared.qn('w:val'), 'single')
    rPr.append(u)
    
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    
    # Add the hyperlink to the paragraph
    paragraph._p.append(hyperlink)
    
    return hyperlink


if __name__ == "__main__":
    pytest.main([__file__, "-v"])
